"""
generate_docx_report.py
-----------------------
Generates a complete, beautifully formatted Word document (.docx) report
for CS559 Assignment 1 - Phase 2: Pond Catchment Analysis Backend.
Embeds the exact screenshots (health.png, postman-results.png, runninglink.png, link.png)
and response data.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

IMG_DIR = Path("/home/ganesh/Desktop/csd/AI-POND PLANNER REPORT/phase2-backend")
OUTPUT_DOCX = Path("/home/ganesh/Desktop/csd/AI-POND-PLANNER/CS559_Assignment1_Phase2_Report.docx")


def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F4F6F8")
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    
    doc.add_paragraph()  # spacing


def main():
    doc = Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Title & Subtitle
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(2)
    t_run = title_p.add_run("CS559 – Computer Systems Design")
    t_run.font.name = 'Arial'
    t_run.font.size = Pt(22)
    t_run.font.bold = True
    t_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)  # Navy Blue

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(16)
    sub_run = sub_p.add_run("Assignment 1 – Phase 2: Pond Catchment Analysis Backend")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(14)
    sub_run.font.bold = True
    sub_run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    # Metadata Card / Table
    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Submitted By:", "Shivam Kushwaha / Ganesh"),
        ("GitHub Repository:", "https://github.com/gnshx/AI-POND-PLANNER"),
        ("Live API URL:", "http://10.1.75.51:4310/analyzeContour")
    ]
    for i, (k, v) in enumerate(meta_data):
        row = meta_table.rows[i]
        c1, c2 = row.cells[0], row.cells[1]
        c1.width = Inches(2.0)
        c2.width = Inches(4.5)
        set_cell_background(c1, "EFF6FF")
        set_cell_background(c2, "F8FAFC")
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(k)
        r1.font.bold = True
        r1.font.size = Pt(10)
        
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(v)
        r2.font.size = Pt(10)
        if "http" in v:
            r2.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            r2.font.underline = True

    doc.add_paragraph()

    def add_h1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    def add_h2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    def add_p(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        p.add_run(text)

    # 1. Introduction
    add_h1("1. Introduction")
    add_p("The objective of this project is to develop a backend system that can analyze a contour map and identify a suitable location for constructing a pond along with the area that naturally drains into that location.")
    add_p("The backend accepts a KML/KMZ contour map through an API request. It processes the contour data, reconstructs the terrain as a Digital Elevation Model (DEM), calculates water flow across the terrain, and identifies a suitable pond location based on the contributing catchment area.")
    add_p("The implementation is designed to work with different contour maps rather than being specifically designed for one sample file.")
    add_p("The main API endpoint is:")
    add_code_block(doc, "POST /analyzeContour")
    add_p("An alternative endpoint is also provided:")
    add_code_block(doc, "POST /findCatchment")

    # 2. System Overview
    add_h1("2. System Overview")
    add_p("The complete processing pipeline consists of four main stages:")
    add_bullet("Contour Map Parsing")
    add_bullet("DEM Reconstruction")
    add_bullet("Flow and Catchment Calculation")
    add_bullet("Pond Location Selection")
    add_p("The overall flow can be represented as:")
    add_code_block(doc, "KML/KMZ File → Contour Data → DEM → Flow Accumulation → Catchment Detection → Pond Location")
    add_p("Each stage is implemented separately so that the system can be extended in future phases.")

    # 3. Contour Map Parsing
    add_h1("3. Contour Map Parsing")
    add_p("The first stage reads the uploaded KML or KMZ file.")
    add_p("The parser extracts the contour information and converts it into elevation samples in the form:")
    add_code_block(doc, "(longitude, latitude, elevation)")
    add_p("Elevation values are obtained primarily from the contour name. Additional parsing methods are included for cases where the elevation is stored inside the KML's ExtendedData or description fields.")
    add_p("This makes the parser more flexible when working with contour maps generated by different tools.")
    add_p("For the sample contour map used in this project:")
    add_bullet("Contour lines: 1,355")
    add_bullet("Contour vertices: 159,113")
    add_bullet("Elevation range: approximately 267 m to 298 m")
    add_bullet("Longitude range: approximately 81.281°E to 81.313°E")
    add_bullet("Latitude range: approximately 21.240°N to 21.264°N")

    # 4. Digital Elevation Model Reconstruction
    add_h1("4. Digital Elevation Model Reconstruction")
    add_p("After extracting the contour points, the system reconstructs the terrain using a regular grid.")
    add_p("The longitude and latitude coordinates are first converted into a local metre-based coordinate system. The elevation values are then interpolated onto the grid using Delaunay triangulation with linear interpolation.")
    add_p("Nearest-neighbour interpolation is used around the edges where linear interpolation does not provide values.")
    add_p("A small amount of Gaussian smoothing is subsequently applied to reduce sharp artifacts created during interpolation.")
    add_p("The grid size is determined from the actual dimensions of the uploaded contour map. Therefore, the system does not depend on a fixed grid specifically designed for the sample input.")
    add_p("For the demonstration input, the generated DEM contains approximately 250,000 cells.")

    # 5. Flow Routing and Catchment Analysis
    add_h1("5. Flow Routing and Catchment Analysis")
    add_p("Once the DEM is generated, the system calculates the direction in which water would flow from each grid cell.")
    add_p("The implementation uses the D8 flow-routing approach. Each cell considers its eight neighbouring cells and selects the neighbouring cell with the steepest downhill slope as its flow direction.")
    add_p("After calculating the flow directions, the system performs flow accumulation. This determines how many upstream cells contribute water to each location.")
    add_p("Locations with high accumulation values represent natural drainage channels, while local low points can represent possible pond locations.")
    add_p("The system then searches for suitable interior sinks and evaluates their contributing areas.")

    # 6. Pond Location Selection
    add_h1("6. Pond Location Selection")
    add_p("The pond location is selected using the calculated terrain and flow information.")
    add_p("The system looks for an interior topographic sink with a sufficiently large contributing area. Areas that exceed the configured main-river accumulation threshold can be excluded so that the selected pond location does not fall directly on a major drainage channel.")
    add_p("The selected location therefore represents an off-stream pond site with a useful upstream catchment.")
    add_p("After selecting the pond point, the flow network is traversed in reverse to identify all cells that drain toward that point.")
    add_p("The resulting set of cells forms the catchment area.")

    # 7. API Implementation
    add_h1("7. API Implementation")
    add_p("The backend is implemented using FastAPI and runs using Uvicorn.")
    add_h2("Main Endpoint")
    add_code_block(doc, "POST /analyzeContour")
    add_p("The endpoint accepts a contour map as multipart form data.")
    add_h2("Request Example")
    add_code_block(doc, 'curl -X POST "http://10.1.75.51:4310/analyzeContour" \\\n  -F "contour_map=@contours_1m.kml"')
    add_p("The uploaded field name is contour_map. The API also supports file as an alias.")
    
    add_h2("Parameters")
    param_table = doc.add_table(rows=6, cols=4)
    param_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Parameter", "Type", "Default", "Description"]
    for j, h in enumerate(headers):
        cell = param_table.cell(0, j)
        set_cell_background(cell, "1E3A8A")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    params_data = [
        ("contour_map", "File", "Required", "KML/KMZ contour map"),
        ("target_cells", "Integer", "250000", "DEM grid resolution"),
        ("min_catchment_fraction", "Float", "0.0001", "Minimum candidate basin size"),
        ("max_river_fraction", "Float", "0.0015", "Main drainage threshold"),
        ("avoid_main_river", "Boolean", "true", "Avoid major drainage channels")
    ]
    for i, row in enumerate(params_data):
        for j, val in enumerate(row):
            cell = param_table.cell(i+1, j)
            set_cell_background(cell, "F8FAFC" if i % 2 == 0 else "FFFFFF")
            p = cell.paragraphs[0]
            p.add_run(val)
            
    doc.add_paragraph()

    # 8. API Response
    add_h1("8. API Response")
    add_p("For the provided sample contour map, the API returned a pond location and catchment information.")
    add_h2("Pond Location")
    add_bullet("Longitude: 81.2886285")
    add_bullet("Latitude: 21.2526051")
    add_bullet("Elevation: 269.01 m")
    
    add_h2("Catchment")
    add_bullet("Area: 49,610 m²")
    add_bullet("Area: 4.961 hectares")
    add_bullet("Cells: 1,450")
    add_bullet("Cell size: 5.849 m")
    add_bullet("Elevation range: 269.01–287.99 m")
    add_bullet("Relief: 18.98 m")
    add_p("The API also returns the catchment boundary as a GeoJSON-style polygon.")
    
    add_h2("DEM Summary")
    add_bullet("Rows: 453")
    add_bullet("Columns: 555")
    add_bullet("Cell size: 5.849 m")
    add_bullet("Minimum elevation: 267.0 m")
    add_bullet("Maximum elevation: 297.89 m")

    # 9. Sample API Response
    add_h1("9. Sample API Response")
    add_p("The complete response generated during testing is stored in docs/response.json.")
    add_p("The response contains the selected pond coordinates, elevation, catchment area, catchment boundary, DEM information, and input contour statistics.")
    add_p("A shortened representation is:")
    short_json = """{
  "pond_location": {
    "longitude": 81.2886285,
    "latitude": 21.2526051,
    "elevation_m": 269.01
  },
  "catchment": {
    "area_m2": 49610.0,
    "area_hectares": 4.961,
    "cell_count": 1450,
    "cell_size_m": 5.849,
    "relief_m": 18.98
  },
  "dem_summary": {
    "grid_rows": 453,
    "grid_cols": 555,
    "cell_size_m": 5.849
  }
}"""
    add_code_block(doc, short_json)

    # 10. API Testing and Demonstration
    add_h1("10. API Testing and Demonstration")
    add_p("The API was tested using the provided contour map and Postman/cURL requests. The following screenshots provide evidence of the implementation and testing.")

    # Helper for adding image figures
    def add_figure(img_input, caption, fig_num):
        if isinstance(img_input, Path):
            img_path = img_input
        else:
            img_path = IMG_DIR / img_input

        if img_path.exists():
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(8)
            p_img.paragraph_format.space_after = Pt(4)
            run = p_img.add_run()
            run.add_picture(str(img_path), width=Inches(6.0))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(12)
            r_cap = p_cap.add_run(f"Figure {fig_num} – {caption}")
            r_cap.font.size = Pt(9.5)
            r_cap.font.italic = True
            r_cap.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
        else:
            add_p(f"[Image file {img_path} not found]")

    # 10.1 Health Check
    add_h2("10.1 Health Check")
    add_p("The /health endpoint was tested to confirm that the deployed backend is running correctly.")
    add_figure("health.png", "API Health Check", 1)
    add_p("The endpoint returned:")
    add_code_block(doc, '{\n  "status": "ok"\n}')
    add_p("This confirms that the deployed FastAPI application is active and responding to requests.")

    # 10.2 Postman API Request and Response
    add_h2("10.2 Postman API Request and Response")
    add_p("The contour analysis endpoint was tested using Postman by uploading the sample KML file.")
    add_figure("postman-results.png", "Postman Analysis Request and Response", 2)
    add_p("The response contains the calculated pond location, catchment area, DEM information, and input contour statistics.")
    add_p("This demonstrates that the backend successfully processes the uploaded contour file and produces the expected analysis results.")

    # 10.3 Running/Deployment Evidence
    add_h2("10.3 Running/Deployment Evidence")
    add_figure("runninglink.png", "Running API Evidence", 3)
    add_p("This screenshot shows the deployed API running and provides evidence that the application is accessible through a live endpoint rather than only running locally.")

    # 10.4 Live API Link
    add_h2("10.4 Live API Link")
    add_figure("link.png", "Live API URL", 4)
    add_p("The live URL can be used to access the deployed backend and its FastAPI documentation.")
    add_bullet("Live API Base URL: http://10.1.75.51:4310/")
    add_bullet("Swagger Documentation: http://10.1.75.51:4310/docs")
    add_bullet("Health Check: http://10.1.75.51:4310/health")

    # 10.5 Terrain Visualization & Catchment Mapping
    add_h2("10.5 Terrain Visualization & Catchment Mapping")
    add_p("The system generates a 3-panel visualization figure illustrating the reconstructed Digital Elevation Model (DEM), flow accumulation channels, and the delineated catchment area along with the selected pond pour point.")
    demo_img_path = Path("/home/ganesh/Desktop/csd/AI-POND-PLANNER/docs/demo_output.png")
    add_figure(demo_img_path, "3-Panel DEM, Flow Accumulation, and Catchment Delineation Visual Map", 5)

    # 10.6 Interactive Swagger API Documentation UI
    add_h2("10.6 Interactive Swagger API Documentation UI")
    add_p("The FastAPI interactive Swagger UI docs interface was tested directly at http://10.1.75.51:4310/docs to verify parameter inputs and multipart form-data specification for contour_map.")
    swagger_img_path = Path("/home/ganesh/Desktop/csd/AI-POND PLANNER REPORT/phase2-backend/swagger-docs.png")
    add_figure(swagger_img_path, "FastAPI Interactive Swagger UI Documentation (/docs)", 6)


    # 11. Project Structure
    add_h1("11. Project Structure")
    add_p("The backend is organized into separate modules according to their responsibilities.")
    tree_str = """AI-POND-PLANNER/
│
├── app/
│   ├── main.py
│   ├── kml_parser.py
│   ├── dem.py
│   ├── hydrology.py
│   └── catchment.py
│
├── scripts/
│   ├── demo_plot.py
│   └── generate_docx_report.py
│
├── docs/
│   ├── demo_output.png
│   ├── swagger-docs.png
│   └── response.json
│
├── requirements.txt
└── README.md"""
    add_code_block(doc, tree_str)
    add_p("The main responsibilities are:")
    add_bullet("main.py – FastAPI routes and request handling")
    add_bullet("kml_parser.py – KML/KMZ contour extraction")
    add_bullet("dem.py – terrain reconstruction and interpolation")
    add_bullet("hydrology.py – D8 flow routing and accumulation")
    add_bullet("catchment.py – pond selection and catchment delineation")
    add_bullet("demo_plot.py – visualization and demonstration output")
    add_p("This separation makes the backend easier to test and modify.")

    # 12. Running the Project Locally
    add_h1("12. Running the Project Locally")
    add_p("The project can be run locally using Python and Uvicorn.")
    add_code_block(doc, "python3 -m venv venv\nsource venv/bin/activate\npip install -r requirements.txt\nuvicorn app.main:app --host 0.0.0.0 --port 4000")
    add_p("After starting the server, the API documentation can be accessed through:")
    add_code_block(doc, "http://localhost:4000/docs")
    add_p("The main analysis endpoint is:")
    add_code_block(doc, "POST /analyzeContour")

    # 13. Technology Stack
    add_h1("13. Technology Stack")
    add_p("The project uses the following technologies:")
    
    tech_table = doc.add_table(rows=10, cols=2)
    tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    cell0 = tech_table.cell(0, 0)
    cell1 = tech_table.cell(0, 1)
    cell0.width = Inches(2.5)
    cell1.width = Inches(4.0)
    set_cell_background(cell0, "1E3A8A")
    set_cell_background(cell1, "1E3A8A")
    
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run("Technology")
    r0.font.bold = True
    r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    p1 = cell1.paragraphs[0]
    r1 = p1.add_run("Purpose")
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    stack_data = [
        ("Python", "Backend implementation"),
        ("FastAPI", "REST API framework"),
        ("Uvicorn", "Application server"),
        ("NumPy", "Numerical and grid operations"),
        ("SciPy", "Delaunay triangulation and interpolation"),
        ("scikit-image", "Catchment boundary extraction"),
        ("KML/KMZ", "Input contour-map format"),
        ("JSON", "API response format"),
        ("Postman", "API testing")
    ]
    for i, (tech, purp) in enumerate(stack_data):
        c_t = tech_table.cell(i+1, 0)
        c_p = tech_table.cell(i+1, 1)
        c_t.width = Inches(2.5)
        c_p.width = Inches(4.0)
        set_cell_background(c_t, "F8FAFC" if i % 2 == 0 else "FFFFFF")
        set_cell_background(c_p, "F8FAFC" if i % 2 == 0 else "FFFFFF")
        c_t.paragraphs[0].add_run(tech)
        c_p.paragraphs[0].add_run(purp)

    doc.add_paragraph()
    add_p("The implementation avoids requiring large GIS frameworks such as GDAL, PostGIS, or external hydrology software. The terrain and hydrology calculations are performed directly using Python numerical libraries.")

    # 14. Repository
    add_h1("14. Repository")
    add_p("The complete source code is available in the GitHub repository:")
    add_bullet("Repository Name: AI-POND-PLANNER")
    add_bullet("GitHub URL: https://github.com/gnshx/AI-POND-PLANNER")
    add_p("The repository contains the backend source code, configuration, demonstration files, and sample API response.")

    # 15. Conclusion
    add_h1("15. Conclusion")
    add_p("The Phase 2 backend successfully implements a contour-based pond catchment analysis system.")
    add_p("The system accepts a KML/KMZ contour map, reconstructs the terrain, calculates D8 water flow and accumulation, identifies a suitable pond location, and determines the corresponding catchment area.")
    add_p("For the provided sample contour map, the system identified a pond location at approximately 81.2886285°E, 21.2526051°N, with an estimated catchment area of 49,610 m² (4.961 hectares).")
    add_p("The API was tested using the deployed backend and Postman, with the results stored in the provided JSON response. The modular implementation also provides a foundation for using different contour maps and extending the system in future phases.")

    # 16. Acknowledgements & AI Assistance Statement
    add_h1("16. Acknowledgements & AI Assistance Statement")
    add_p("Taken help from Gemini AI in designing system architecture, refining hydrological flow-routing algorithms, compiling report documentation, and creating comprehensive unit and integration test suites.")

    doc.save(OUTPUT_DOCX)
    print(f"Report successfully saved to {OUTPUT_DOCX}")

    # Convert DOCX to PDF using LibreOffice
    import subprocess
    output_pdf = OUTPUT_DOCX.with_suffix('.pdf')
    print("Converting DOCX to PDF via LibreOffice...")
    cmd = ["libreoffice", "--headless", "--convert-to", "pdf", str(OUTPUT_DOCX), "--outdir", str(OUTPUT_DOCX.parent)]
    subprocess.run(cmd, check=True)
    print(f"PDF successfully generated at {output_pdf}")


if __name__ == "__main__":
    main()
